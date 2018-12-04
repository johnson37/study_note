import docx
from tkinter import *
import tkinter.filedialog


pre_a = []
pre_b = []
pre_c = []
pre_d = []
pre_e = []
pre_f = []
pre_g = []
pre_h = []
pre_i = []
pre_j = []
pre_k = []
pre_l = []
pre_m = []
pre_n = []
pre_o = []
pre_p = []
pre_q = []
pre_r = []
pre_s = []
pre_t = []
pre_u = []
pre_v = []
pre_w = []
pre_x = []
pre_y = []
pre_z = []
mydic = {}
mydic['a']=pre_a
mydic['b']=pre_b
mydic['c']=pre_c
mydic['d']=pre_d
mydic['e']=pre_e
mydic['f']=pre_f
mydic['g']=pre_g
mydic['h']=pre_h
mydic['i']=pre_i
mydic['j']=pre_j
mydic['k']=pre_k
mydic['l']=pre_l
mydic['m']=pre_m
mydic['n']=pre_n
mydic['o']=pre_o
mydic['p']=pre_p
mydic['q']=pre_q
mydic['r']=pre_r
mydic['s']=pre_s
mydic['t']=pre_t
mydic['u']=pre_u
mydic['v']=pre_v
mydic['w']=pre_w
mydic['x']=pre_x
mydic['y']=pre_y
mydic['z']=pre_z

#file
source = ""
file_list = []
def readStandard():
    f = open("D:\\standard.txt","r")
    standard_str = f.read()
    #print (standard_str)
    standard = standard_str.split('\n')
    for standard_word in standard:
        #print (standard_word)
        mydic[standard_word[0]].append(standard_word)
    #print(pre_c)
    print("Finish")
    return "helloworld"

def searchElement(element):
    try:
        #Change all charater to lower
        new_element = element.lower()
        # Clear some special character
        new_element = new_element.strip('\'”“",.?:()')
        # Handle 's issue
        if (len(new_element)> 2):
            #print (new_element[-2:])
            if (new_element[-2:] ==  "’s"):
                new_element = new_element[:-2]
            elif new_element[-2:] == "'s":
                new_element = new_element[:-2]
        mydic[new_element[0]].index(new_element)    
    except:
        return False
    else:
        return True
def checkWord(file):
    readStandard()
    for para in file.paragraphs:
        # Get all the text
        text = para.text
        # Clear all the paragraph
        para.clear()
        # Split the words in this paragraph through space
        elements = text.split()
        for element in elements:
            run = para.add_run(element)
            if searchElement(element) ==  False:
                run.bold = True
                run.underline = True
                run.font.color.rgb = docx.shared.RGBColor(0xff, 0x00, 0x00)
            para.add_run(" ")
        #print (element)

def getFile():
    source = tkinter.filedialog.askopenfilename()
    if source == ' ':
        return
    file = docx.Document(source)
    checkWord(file)
    global file_list
    file_list.append(file)
    print (file_list)
    return

def saveFile():
    fname = tkinter.filedialog.asksaveasfilename(title=u'保存文件')
    if fname == ' ':
        return
    if len(fname) == 0:
        return

    fname=fname+".docx"
    if len(file_list) > 0 :
        file_list[0].save(fname)
        print ("Save File")
        root.destroy()
    return

root = Tk()
root.geometry("800x400")
btn = Button(root,text="Choose the source file",command=getFile)
btn.pack()

save_btn = Button(root,text="Save the trans file",command=saveFile)
save_btn.pack()

#file.save("D:\\51.docx")
root.mainloop()




